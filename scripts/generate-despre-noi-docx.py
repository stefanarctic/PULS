# -*- coding: utf-8 -*-
"""Generează docs/Despre-noi-PULS.docx din același text ca în src/components/About.jsx (UTF-8)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "docs" / "Despre-noi-PULS.docx"

PARAGRAPHS = [
    "P.U.L.S. este o platformă educațională pentru fizică care combină simulări interactive, probleme și grile, resurse teoretice pe mai multe capitole și instrumente digitale menite să facă abstractul vizibil. Ne-am propus să mergem dincolo de „pagină cu formule”: vrem un loc unde poți exersa pentru școală sau bac, să urmărești progresul și să explorezi fenomene, de la oscilații și unde la mecanică, termodinamică, optică sau astronomie, într-un mod modern, clar și captivant.",
    "Povestea noastră a început cu o întrebare simplă: cum putem face fenomenele fizice, în special cele oscilatorii, să prindă viață și să devină mai ușor de înțeles? Noi, o echipă de elevi pasionați de știință, am simțit mereu că, dincolo de formule și definiții, există o lume fascinantă, plină de ritm, mișcare și conexiuni surprinzătoare cu natura și tehnologia. De la acel punct de plecare, P.U.L.S. a crescut într-o platformă mai largă: simulări și experimente virtuale, colecții de probleme și variante de tip bac, grile pentru antrenament rapid, materiale și lecții grupate pe tematici, plus spațiu pentru profil și statistici ca să îți vezi evoluția. În spate stă aceeași curiozitate: fizica nu e doar simboluri pe tablă, ci mișcare, ritm și legături cu lumea reală.",
    "Am îmbinat pasiunea pentru fizică cu dezvoltare web modernă și instrumente de simulare (inclusiv medii precum Unity acolo unde experiența cere grafică și interacțiune avansată), astfel încât să ai atât „laboratorul” în browser, cât și exercițiile și resursele într-un singur loc. Am inclus atât experimente virtuale, cât și secvențe filmate în laborator, acolo unde ajută la înțelegere. Ne dorim ca oricine învață fizică să găsească aici nu doar o simulare izolată, ci un parcurs: de la idee, la vizualizare, la practică și repetare.",
    "Nu a fost un drum ușor. Am petrecut ore întregi studiind, programând și testând fiecare simulare, cu sprijinul necondiționat al mentorului nostru, doamna profesoară Bebu Ioana Bianka, și al colaboratorului nostru, fizicianul Bebu Ion. Am învățat că, dincolo de cifre și ecuații, fizica este o poveste despre ordine și haos, despre cum universul dansează pe ritmuri nevăzute, dar perfect regulate.",
    "P.U.L.S. nu este doar un site, ci o fereastră către un univers fascinant. Este dovada că știința poate fi frumoasă, că poate inspira și că, prin curiozitate și muncă, putem face ca lucrurile complicate să devină accesibile și uimitoare.",
    "Așa că, dacă vrei să vezi cum se mișcă lumea într-un ritm nevăzut, te invităm să explorezi platforma. Descoperă, experimentează și lasă-te purtat de pulsul fascinant al științei!",
]

MOTTO = "Descoperă, experimentează și lasă-te purtat de PULS-ul fascinant al științei!"


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_heading("Despre noi", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for text in PARAGRAPHS:
        doc.add_paragraph(text)

    motto_p = doc.add_paragraph()
    motto_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = motto_p.add_run(MOTTO)
    run.bold = True
    run.italic = True

    try:
        doc.save(OUT)
        print(f"Saved: {OUT}")
    except PermissionError:
        alt = OUT.with_name(f"{OUT.stem}-nou{OUT.suffix}")
        doc.save(alt)
        print(f"Original locked; saved: {alt}")


if __name__ == "__main__":
    main()
